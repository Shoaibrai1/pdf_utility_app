import streamlit as st
from PIL import Image
import os
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from pdf2image import convert_from_path
import fitz  # PyMuPDF

st.set_page_config(page_title="PDF Utility App", layout="centered")

# Custom CSS for mobile responsiveness
st.markdown("""
    <style>
    body { zoom: 0.8; }
    .block-container {
        padding-top: 1rem;
    }
    </style>
""", unsafe_allow_html=True)

st.title("üìÑ PDF Utility Web App")

# Sidebar menu
task = st.sidebar.selectbox("Choose Task", [
    "Images to PDF",
    "PDF to JPG",
    "DOCX to PDF",
    "PDF to DOCX",
    "Merge PDFs",
    "Split PDF",
    "Protect PDF"
])

# 1. Images to PDF
if task == "Images to PDF":
    st.header("üñºÔ∏è Convert Images to PDF")
    images = st.file_uploader("Upload multiple images", type=['jpg', 'jpeg', 'png'], accept_multiple_files=True)
    if images:
        img_list = []
        for image_file in images:
            img = Image.open(image_file).convert('RGB')
            img_list.append(img)
        if st.button("Convert to PDF"):
            pdf_bytes = io.BytesIO()
            img_list[0].save(pdf_bytes, format='PDF', save_all=True, append_images=img_list[1:])
            st.download_button("üì• Download PDF", data=pdf_bytes.getvalue(), file_name="images_to_pdf.pdf")

# 2. PDF Edit (Add Text)
elif task == "Edit PDF (Add Text)":
    st.header("‚úèÔ∏è Edit PDF - Add Custom Text to Pages")
    pdf_file = st.file_uploader("Upload PDF", type="pdf")
    if pdf_file:
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        text_inputs = []
        st.write("Enter text to add on each page:")
        for i, page in enumerate(doc):
            text = st.text_input(f"Text for Page {i+1}", key=f"text_page_{i}")
            text_inputs.append(text)
        if st.button("Apply and Download Edited PDF"):
            for i, page in enumerate(doc):
                if text_inputs[i]:
                    page.insert_text((50, page.rect.height - 50), text_inputs[i], fontsize=12, color=(0, 0, 0))
            output = io.BytesIO()
            doc.save(output)
            st.download_button("üì• Download Edited PDF", data=output.getvalue(), file_name="edited_custom.pdf")

# 3. PDF to JPG
elif task == "PDF to JPG":
    st.header("üßæ Convert PDF to JPG")
    pdf_file = st.file_uploader("Upload PDF", type="pdf")
    if pdf_file:
        with open("temp.pdf", "wb") as f:
            f.write(pdf_file.read())
        images = convert_from_path("temp.pdf")
        for i, image in enumerate(images):
            img_buffer = io.BytesIO()
            image.save(img_buffer, format="JPEG")
            st.image(image, caption=f"Page {i+1}")
            st.download_button(f"Download Page {i+1}", img_buffer.getvalue(), file_name=f"page_{i+1}.jpg")

# 4. DOCX to PDF
elif task == "DOCX to PDF":
    st.header("üìÑ Convert DOCX to PDF")
    docx_file = st.file_uploader("Upload DOCX", type="docx")
    if docx_file:
        doc = Document(docx_file)
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        textobject = c.beginText(40, height - 40)
        for para in doc.paragraphs:
            textobject.textLine(para.text)
        c.drawText(textobject)
        c.save()
        buffer.seek(0)
        st.download_button("üì• Download PDF", data=buffer.getvalue(), file_name="docx_to_pdf.pdf")
# 6. Split PDF
elif task == "Split PDF":
    st.header("‚úÇÔ∏è Split PDF by Page Range")
    pdf_file = st.file_uploader("Upload PDF", type="pdf")
    start_page = st.number_input("Start Page (1-based)", min_value=1, value=1)
    end_page = st.number_input("End Page", min_value=start_page, value=start_page)
    if pdf_file and st.button("Split PDF"):
        from PyPDF2 import PdfReader, PdfWriter
        reader = PdfReader(pdf_file)
        writer = PdfWriter()
        total_pages = len(reader.pages)
        start = max(0, start_page - 1)
        end = min(end_page, total_pages)
        for i in range(start, end):
            writer.add_page(reader.pages[i])
        split_output = io.BytesIO()
        writer.write(split_output)
        st.download_button("üì• Download Split PDF", data=split_output.getvalue(), file_name="split.pdf")
# 7. Protect PDF
elif task == "Protect PDF":
    st.header("üîê Protect PDF with Password")
    pdf_file = st.file_uploader("Upload PDF", type="pdf")
    password = st.text_input("Enter Password to Protect PDF")
    if pdf_file and password and st.button("Protect PDF"):
        from PyPDF2 import PdfReader, PdfWriter
        reader = PdfReader(pdf_file)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(password)
        protected_output = io.BytesIO()
        writer.write(protected_output)
        st.download_button("üì• Download Protected PDF", data=protected_output.getvalue(), file_name="protected.pdf")

# 5. PDF to DOCX
elif task == "PDF to DOCX":
    st.header("üîÅ Convert PDF to DOCX")
    pdf_file = st.file_uploader("Upload PDF", type="pdf")
    if pdf_file:
        docx_doc = Document()
        doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
        for i, page in enumerate(doc):
            text = page.get_text()
            docx_doc.add_paragraph(f"--- Page {i+1} ---")
            docx_doc.add_paragraph(text)
        buffer = io.BytesIO()
        docx_doc.save(buffer)
        buffer.seek(0)
        st.download_button("üì• Download DOCX", data=buffer.getvalue(), file_name="converted.docx")
